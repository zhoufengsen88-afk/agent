from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = next((ROOT / "word").glob("*LangChain*.docx"))
ASSET_DIR = ROOT / "word" / "generated_assets"
ASSET_DIR.mkdir(exist_ok=True)

FONT = Path(r"C:\Windows\Fonts\msyh.ttc")
BOLD_FONT = Path(r"C:\Windows\Fonts\msyhbd.ttc")
MONO_FONT = Path(r"C:\Windows\Fonts\consola.ttf")


def font(size, bold=False, mono=False):
    path = MONO_FONT if mono and MONO_FONT.exists() else (BOLD_FONT if bold else FONT)
    return ImageFont.truetype(str(path), size)


def wrap_text(draw, text, fnt, width):
    lines = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if draw.textlength(test, font=fnt) <= width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        lines.append(line)
    return lines


def rounded_box(draw, box, text, fill, outline="#355070", title=False):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=14, fill=fill, outline=outline, width=2)
    fnt = font(22 if title else 18, bold=title)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 24)
    total_h = len(lines) * (fnt.size + 8)
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line in lines:
        w = draw.textlength(line, font=fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill="#1f2933")
        y += fnt.size + 8


def arrow(draw, start, end, color="#2f4858"):
    draw.line([start, end], fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    import math

    ang = math.atan2(y2 - y1, x2 - x1)
    for a in (ang + 2.6, ang - 2.6):
        draw.line([end, (x2 + 14 * math.cos(a), y2 + 14 * math.sin(a))], fill=color, width=3)


def save_architecture():
    img = Image.new("RGB", (1500, 950), "#f7fafc")
    d = ImageDraw.Draw(img)
    d.text((520, 35), "基于 LangChain 的智能扫地机器人客服系统总体架构", font=font(30, True), fill="#102a43")
    boxes = [
        ((80, 150, 330, 280), "表示层\nStreamlit 聊天界面\n用户输入 / 流式输出", "#d9f0ff"),
        ((430, 150, 700, 280), "智能体编排层\nReactAgent\ncreate_agent / stream", "#e6fffa"),
        ((800, 150, 1120, 280), "工具与中间件层\nTool Calling\n日志监控 / 提示词切换", "#fff7d6"),
        ((1180, 150, 1430, 280), "模型服务层\n通义千问 qwen3-max\nDashScope Embedding", "#ffe4e6"),
        ((430, 430, 700, 570), "RAG 检索层\nPromptTemplate + Retriever\n资料召回与总结", "#e8f5e9"),
        ((800, 430, 1120, 570), "数据资源层\nChroma 向量库\n知识文件 / 外部记录 CSV", "#efe9ff"),
        ((80, 430, 330, 570), "配置与提示词层\nYAML 配置\nmain/rag/report Prompt", "#f0f4f8"),
        ((430, 700, 700, 830), "运行支撑层\nutils/logger/path/file\n日志脱敏与路径管理", "#fdebd0"),
    ]
    for box, text, fill in boxes:
        rounded_box(d, box, text, fill, title=False)
    for s, e in [
        ((330, 215), (430, 215)),
        ((700, 215), (800, 215)),
        ((1120, 215), (1180, 215)),
        ((960, 280), (960, 430)),
        ((800, 500), (700, 500)),
        ((565, 430), (565, 280)),
        ((330, 500), (430, 500)),
        ((565, 700), (565, 570)),
    ]:
        arrow(d, s, e)
    d.text((610, 890), "图4-1 系统总体架构图", font=font(22, True), fill="#334e68")
    path = ASSET_DIR / "system_architecture.png"
    img.save(path)
    return path


def save_flow():
    img = Image.new("RGB", (1400, 1050), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((530, 30), "系统工作流程图", font=font(32, True), fill="#102a43")
    steps = [
        ((550, 110, 850, 185), "用户输入问题", "#d9f0ff"),
        ((550, 230, 850, 305), "Streamlit 接收并提交", "#d9f0ff"),
        ((520, 350, 880, 425), "ReactAgent 组织消息并调用 LangChain Agent", "#e6fffa"),
        ((520, 470, 880, 545), "模型结合主提示词判断意图", "#e6fffa"),
        ((120, 620, 430, 710), "知识咨询\n调用 rag_summarize\n检索 Chroma 后生成答案", "#e8f5e9"),
        ((545, 620, 855, 710), "报告生成\n获取用户ID/月度\n触发 report=True", "#fff7d6"),
        ((970, 620, 1280, 710), "环境辅助\n获取位置/天气\n补充场景参数", "#ffe4e6"),
        ((545, 790, 855, 870), "中间件记录工具调用\n动态切换提示词", "#efe9ff"),
        ((545, 930, 855, 1000), "前端流式展示回答", "#d9f0ff"),
    ]
    centers = {}
    for box, text, fill in steps:
        rounded_box(d, box, text, fill)
        centers[text.split("\n")[0]] = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
    for s, e in [
        ((700, 185), (700, 230)),
        ((700, 305), (700, 350)),
        ((700, 425), (700, 470)),
        ((620, 545), (275, 620)),
        ((700, 545), (700, 620)),
        ((780, 545), (1125, 620)),
        ((275, 710), (620, 790)),
        ((700, 710), (700, 790)),
        ((1125, 710), (780, 790)),
        ((700, 870), (700, 930)),
    ]:
        arrow(d, s, e)
    d.text((610, 1012), "图4-2 系统工作流程图", font=font(22, True), fill="#334e68")
    path = ASSET_DIR / "system_flow.png"
    img.save(path)
    return path


def save_import_flow():
    img = Image.new("RGB", (1400, 650), "#f8fafc")
    d = ImageDraw.Draw(img)
    d.text((480, 35), "知识库构建与检索流程", font=font(32, True), fill="#102a43")
    labels = ["TXT/PDF 知识文件", "文件遍历与 MD5 去重", "文档加载", "中文语义切分", "DashScope 向量化", "写入 Chroma", "用户问题相似度检索"]
    x = 50
    last = None
    for i, label in enumerate(labels):
        box = (x, 210, x + 170, 340)
        rounded_box(d, box, label, ["#d9f0ff", "#e6fffa", "#fff7d6", "#e8f5e9", "#ffe4e6", "#efe9ff", "#fdebd0"][i])
        if last:
            arrow(d, (last[2], 275), (box[0], 275))
        last = box
        x += 195
    d.text((560, 540), "图4-3 知识库构建与检索流程图", font=font(22, True), fill="#334e68")
    path = ASSET_DIR / "knowledge_import_flow.png"
    img.save(path)
    return path


def save_effect(name, title, user, assistant, footer):
    img = Image.new("RGB", (1200, 720), "#eef2f6")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((80, 55, 1120, 650), radius=22, fill="#ffffff", outline="#cbd5e1", width=2)
    d.text((120, 95), "智扫通机器人智能客服", font=font(32, True), fill="#102a43")
    d.line((120, 145, 1080, 145), fill="#d9e2ec", width=2)
    d.rounded_rectangle((120, 180, 610, 250), radius=18, fill="#f1f5f9")
    d.text((145, 200), "你好，我是智扫通机器人智能客服，请问有什么可以帮助你？", font=font(20), fill="#334155")
    d.rounded_rectangle((520, 285, 1080, 365), radius=18, fill="#dbeafe")
    for i, line in enumerate(wrap_text(d, user, font(21), 510)[:2]):
        d.text((545, 305 + i * 28), line, font=font(21), fill="#1e3a8a")
    d.rounded_rectangle((120, 400, 1010, 560), radius=18, fill="#f8fafc", outline="#e2e8f0")
    d.text((145, 425), title, font=font(22, True), fill="#0f172a")
    for i, line in enumerate(wrap_text(d, assistant, font(19), 820)[:4]):
        d.text((145, 462 + i * 29), line, font=font(19), fill="#334155")
    d.rounded_rectangle((120, 590, 1080, 630), radius=18, fill="#f8fafc", outline="#cbd5e1")
    d.text((145, 600), footer, font=font(18), fill="#64748b")
    path = ASSET_DIR / name
    img.save(path)
    return path


def make_assets():
    return {
        "arch": save_architecture(),
        "flow": save_flow(),
        "import": save_import_flow(),
        "ui": save_effect(
            "effect_ui.png",
            "功能效果：前端对话交互",
            "滤网多久清洗一次？",
            "建议每周清理滤网表面灰尘，并根据使用频率定期更换。若家中宠物毛发较多，应缩短清理周期，避免吸力下降。",
            "请输入您的问题",
        ),
        "qa": save_effect(
            "effect_qa.png",
            "功能效果：知识库问答",
            "小户型适合什么扫地机器人？",
            "小户型可优先选择体积较小、建图稳定、支持分区清扫和自动回充的机型。若地面杂物较多，应关注避障能力。",
            "RAG 检索：选购指南 / 扫地机器人100问",
        ),
        "report": save_effect(
            "effect_report.png",
            "功能效果：个性化使用报告",
            "帮我生成本月扫地机器人使用报告",
            "系统先获取用户ID和月份，再读取外部记录，生成清洁效率、耗材状态、使用对比与保养建议，形成结构化报告。",
            "报告模式：fill_context_for_report -> fetch_external_data",
        ),
        "log": save_effect(
            "effect_log.png",
            "功能效果：日志与监控",
            "我所在城市适合拖地模式吗？",
            "中间件记录模型调用、工具名称和参数，并在日志中过滤密钥、手机号、邮箱等敏感信息，便于调试和追踪。",
            "agent_YYYYMMDD.log：tool monitor / before_model",
        ),
    }


def insert_paragraph_after(paragraph, text="", style=None, align=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if align is not None:
        p.alignment = align
    if text:
        p.add_run(text)
    return p


def insert_picture_after(paragraph, image_path, width=5.8, caption=None):
    p = insert_paragraph_after(paragraph, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        p = insert_paragraph_after(p, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def insert_code_after(paragraph, title, code):
    p = insert_paragraph_after(paragraph, title)
    for run in p.runs:
        run.bold = True
    p = insert_paragraph_after(p, code.strip())
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def find_para(doc, exact):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(f"找不到段落：{exact}")


def add_content():
    assets = make_assets()
    doc = Document(str(DOCX_PATH))

    p = find_para(doc, "4.2 系统总体架构")
    p = insert_picture_after(p, assets["arch"], width=6.1, caption="图4-1 系统总体架构图")
    p = insert_paragraph_after(
        p,
        "如图4-1所示，系统由表示层、智能体编排层、工具与中间件层、RAG检索层、模型服务层、数据资源层、配置与提示词层以及运行支撑层组成。表示层负责用户交互，ReactAgent负责统一调度模型与工具，RAG模块负责从本地知识库召回资料，工具与中间件层负责业务能力扩展、日志监控和报告模式切换。该架构将页面、智能体、知识库、外部数据和运行监控分离，降低了模块耦合度。",
    )

    p = find_para(doc, "4.3 系统工作流程设计")
    p = insert_picture_after(p, assets["flow"], width=6.0, caption="图4-2 系统工作流程图")
    p = insert_paragraph_after(
        p,
        "图4-2展示了系统从用户输入到答案输出的完整链路。普通知识咨询会进入RAG检索路径，报告生成需求会进入用户标识、月份和外部记录读取路径，环境适配类问题则可调用位置和天气工具补充上下文。中间件贯穿工具调用与模型调用过程，使系统在运行时具备可观测性和流程切换能力。",
    )

    p = find_para(doc, "4.4 知识库与数据设计")
    p = insert_picture_after(p, assets["import"], width=6.0, caption="图4-3 知识库构建与检索流程图")
    p = insert_paragraph_after(
        p,
        "知识入库阶段重点解决数据来源、重复入库和中文语义切分问题。系统读取data目录下的txt和pdf文件，通过MD5记录避免重复处理，再利用RecursiveCharacterTextSplitter按中文标点和长度参数切分文本，最后调用向量模型生成嵌入并写入Chroma。问答阶段则直接使用检索器召回相似片段，减少模型脱离资料自由生成的风险。",
    )

    p = find_para(doc, "5.2 前端交互模块实现")
    p = insert_picture_after(p, assets["ui"], width=5.8, caption="图5-1 前端对话功能效果图")
    p = insert_code_after(
        p,
        "核心代码5-1 前端消息接收与流式输出：",
        dedent(
            """
            prompt = st.chat_input()
            if prompt:
                st.chat_message("user").write(prompt)
                st.session_state["message"].append({"role": "user", "content": prompt})
                response_messages = []
                with st.spinner("智能客服思考中..."):
                    res_stream = st.session_state["agent"].execute_stream(prompt)
                    st.chat_message("assistant").write_stream(capture(res_stream, response_messages))
                    st.session_state["message"].append({"role": "assistant", "content": response_messages[-1]})
                    st.rerun()
            """
        ),
    )

    p = find_para(doc, "5.3 智能体构建实现")
    p = insert_code_after(
        p,
        "核心代码5-2 LangChain智能体创建与流式执行：",
        dedent(
            """
            self.agent = create_agent(
                model=chat_model,
                system_prompt=load_system_prompt(),
                tools=[rag_summarize, get_weather, get_user_location, get_user_id,
                       get_current_month, fetch_external_data, fill_context_for_report],
                middleware=[monitor_tool, log_before_model, report_prompt_switch],
            )

            for chunk in self.agent.stream(input_dict, stream_mode="values", context={"report": False}):
                latest_message = chunk["messages"][-1]
                if latest_message.content:
                    yield latest_message.content.strip() + "\\n"
            """
        ),
    )

    p = find_para(doc, "5.4 工具模块实现")
    p = insert_picture_after(p, assets["qa"], width=5.8, caption="图5-2 知识库问答功能效果图")
    p = insert_picture_after(p, assets["report"], width=5.8, caption="图5-3 个性化报告生成功能效果图")
    p = insert_code_after(
        p,
        "核心代码5-3 工具函数注册与外部记录读取：",
        dedent(
            """
            @tool(description="从向量存储中检索参考资料")
            def rag_summarize(query: str) -> str:
                return rag.rag_summarize(query)

            @tool(description="从外部系统中获取指定用户在指定月份的使用记录")
            def fetch_external_data(user_id: str, month: str) -> str:
                generrate_external_data()
                try:
                    return external_data[user_id][month]
                except KeyError:
                    logger.warning(f"[fetch_external_data]未能检索到用户：{user_id}在{month}使用记录数据")
                    return ""
            """
        ),
    )

    p = find_para(doc, "5.5 中间件实现")
    p = insert_picture_after(p, assets["log"], width=5.8, caption="图5-4 日志监控功能效果图")
    p = insert_code_after(
        p,
        "核心代码5-4 工具监控与报告提示词切换：",
        dedent(
            """
            @wrap_tool_call
            def monitor_tool(request: ToolCallRequest, handler):
                logger.info(f"[tool monitor]执行工具: {request.tool_call['name']}")
                result = handler(request)
                if request.tool_call['name'] == 'fill_context_for_report':
                    request.runtime.context["report"] = True
                return result

            @dynamic_prompt
            def report_prompt_switch(request: ModelRequest) -> str:
                is_report = request.runtime.context.get("report", False)
                return load_report_prompt() if is_report else load_system_prompt()
            """
        ),
    )

    p = find_para(doc, "5.6 检索增强生成模块实现")
    p = insert_code_after(
        p,
        "核心代码5-5 RAG检索与答案生成链：",
        dedent(
            """
            self.retriever = self.vector_store.get_retriever()
            self.prompt_template = PromptTemplate.from_template(load_rag_prompt())
            self.chain = self.prompt_template | print_prompt | self.model | StrOutputParser()

            def rag_summarize(self, query: str) -> str:
                context_docs = self.retriever_docs(query)
                context = ""
                for counter, doc in enumerate(context_docs, 1):
                    context += f"【参考资料{counter}：参考资料{doc.page_content} | 参考元数据：{doc.metadata}\\n】"
                return self.chain.invoke({"input": query, "context": context})
            """
        ),
    )
    p = insert_code_after(
        p,
        "核心代码5-6 知识文件切分与向量库写入：",
        dedent(
            """
            self.spliter = RecursiveCharacterTextSplitter(
                chunk_size=chroma_conf["chunk_size"],
                chunk_overlap=chroma_conf["chunk_overlap"],
                separators=chroma_conf["separators"],
                length_function=len,
            )
            split_document = self.spliter.split_documents(documents)
            self.vector_store.add_documents(split_document)
            """
        ),
    )

    p = find_para(doc, "5.7 模型与配置管理实现")
    p = insert_code_after(
        p,
        "核心代码5-7 模型工厂与配置化创建：",
        dedent(
            """
            class ChatModelFactory(BaseModelFactory):
                def generator(self):
                    return ChatTongyi(model=rag_conf["chat_model_name"])

            class EmbeddingsFactory(BaseModelFactory):
                def generator(self):
                    return DashScopeEmbeddings(model=rag_conf["embedding_model_name"])

            chat_model = ChatModelFactory().generator()
            embed_model = EmbeddingsFactory().generator()
            """
        ),
    )

    p = find_para(doc, "5.8 日志与安全实现")
    p = insert_code_after(
        p,
        "核心代码5-8 日志敏感信息脱敏：",
        dedent(
            """
            def mask_sensitive_data(text: str) -> str:
                text = re.sub(r"sk-\\w+", "sk-******", text)
                text = re.sub(r"1[3-9]\\d{9}", "1**********", text)
                text = re.sub(r"(\\w+)@(\\w+)\\.(\\w+)", r"\\1****@\\2.\\3", text)
                text = re.sub(r"(password|key|secret)=[^& ]+", r"\\1=******", text)
                return text
            """
        ),
    )

    p = find_para(doc, "6.3.1 前端交互测试")
    p = insert_paragraph_after(
        p,
        "测试效果如图5-1所示。页面能够完成欢迎语展示、用户消息渲染、等待状态提示和助手消息流式输出，说明前端主流程已经形成闭环。",
    )
    p = find_para(doc, "6.3.2 知识问答测试")
    p = insert_paragraph_after(
        p,
        "测试效果如图5-2所示。系统能够围绕用户问题调用知识检索工具，并将选购、维护、故障等领域资料转化为自然语言答案。",
    )
    p = find_para(doc, "6.3.3 报告流程测试")
    p = insert_paragraph_after(
        p,
        "测试效果如图5-3所示。报告场景不是直接由模型自由生成，而是通过用户ID、月份、上下文注入和外部记录查询组成受控链路，更符合客服系统对个性化数据的处理要求。",
    )
    p = find_para(doc, "6.3.4 日志与监控测试")
    p = insert_paragraph_after(
        p,
        "测试效果如图5-4所示。日志功能能够记录工具调用与模型调用状态，配合脱敏过滤器后可在调试便利性和数据安全之间取得平衡。",
    )

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    add_content()
    print(f"已修改：{DOCX_PATH}")
    print(f"插图目录：{ASSET_DIR}")
